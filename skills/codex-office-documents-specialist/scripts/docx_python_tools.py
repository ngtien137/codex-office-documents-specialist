#!/usr/bin/env python3
"""Shared python-docx helpers for Word document generation and template filling."""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, Iterable, List, Optional, Tuple

import requests
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

PLACEHOLDER_PATTERN = re.compile(r"\{\{\{?([a-zA-Z_][a-zA-Z0-9_]*)\}?\}\}")


def add_hyperlink(paragraph: Paragraph, text: str, url: str, color: str = "0000FF", underline: bool = True) -> None:
    """Insert an external hyperlink into a paragraph."""
    try:
        part = paragraph.part
        relationship_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        if underline:
            underline_node = OxmlElement("w:u")
            underline_node.set(qn("w:val"), "single")
            run_properties.append(underline_node)

        if color:
            color_node = OxmlElement("w:color")
            color_node.set(qn("w:val"), color)
            run_properties.append(color_node)

        run.append(run_properties)

        text_node = OxmlElement("w:t")
        text_node.text = text
        text_node.set(qn("xml:space"), "preserve")
        run.append(text_node)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)
    except Exception as exc:
        logger.warning("Failed to create hyperlink for %s (%s): %s", text, url, exc)
        paragraph.add_run(text)


def _handle_escapes(text: str, escape_context: Dict[str, Any]) -> str:
    def replace(match: re.Match[str]) -> str:
        placeholder = chr(0xE000 + escape_context["counter"])
        escape_context["map"][placeholder] = match.group(1)
        escape_context["counter"] += 1
        return placeholder

    return re.sub(r"\\(.)", replace, text)


def _restore_escapes(text: str, escape_context: Optional[Dict[str, Any]]) -> str:
    if not escape_context:
        return text
    restored = text
    for placeholder, original in escape_context["map"].items():
        restored = restored.replace(placeholder, original)
    return restored


def _apply_inherited_formatting(run: Any, bold: bool = False, italic: bool = False) -> None:
    if bold:
        run.bold = True
    if italic:
        run.italic = True


INLINE_FORMAT_RE = re.compile(
    r"(\*{3}(?:[^*]|\*(?!\*{2}))+\*{3}"
    r"|\*\*(?:[^*]|\*(?!\*))+\*\*"
    r"|~~.+?~~"
    r"|__(?!_).+?__"
    r"|\*(?:[^*]|\*\*[^*]+\*\*)+\*"
    r"|`[^`]+`"
    r"|\[[^\]]*\]\([^)]*\))"
)


def parse_inline_formatting(text: str, paragraph: Paragraph, bold: bool = False, italic: bool = False) -> None:
    """Parse inline markdown into runs on an existing paragraph."""
    escape_context = {"map": {}, "counter": 0}
    text = _handle_escapes(text, escape_context)

    line_parts = text.split("  \n")
    for index, line_part in enumerate(line_parts):
        if not line_part and index == len(line_parts) - 1:
            continue
        _parse_formatting_segment(line_part, paragraph, bold=bold, italic=italic, escape_context=escape_context)
        if index < len(line_parts) - 1:
            paragraph.add_run().add_break()


def _parse_formatting_segment(
    text: str,
    paragraph: Paragraph,
    bold: bool = False,
    italic: bool = False,
    escape_context: Optional[Dict[str, Any]] = None,
) -> None:
    for part in INLINE_FORMAT_RE.split(text):
        if not part:
            continue
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            _parse_formatting_segment(part[3:-3], paragraph, bold=True, italic=True, escape_context=escape_context)
        elif part.startswith("**") and part.endswith("**"):
            _parse_formatting_segment(part[2:-2], paragraph, bold=True, italic=italic, escape_context=escape_context)
        elif part.startswith("~~") and part.endswith("~~"):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_context))
            run.font.strike = True
            _apply_inherited_formatting(run, bold=bold, italic=italic)
        elif part.startswith("__") and part.endswith("__") and not part.startswith("___"):
            run = paragraph.add_run(_restore_escapes(part[2:-2], escape_context))
            run.font.underline = True
            _apply_inherited_formatting(run, bold=bold, italic=italic)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            _parse_formatting_segment(part[1:-1], paragraph, bold=bold, italic=True, escape_context=escape_context)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(_restore_escapes(part[1:-1], escape_context))
            run.font.name = "Courier New"
            _apply_inherited_formatting(run, bold=bold, italic=italic)
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            match = re.match(r"\[(.*?)]\((.*?)\)", part)
            if match:
                add_hyperlink(paragraph, _restore_escapes(match.group(1), escape_context), _restore_escapes(match.group(2), escape_context))
        else:
            run = paragraph.add_run(_restore_escapes(part, escape_context))
            _apply_inherited_formatting(run, bold=bold, italic=italic)


def snapshot_run_format(run: Any) -> Dict[str, Any]:
    """Capture enough run formatting to preserve template appearance."""
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.font.underline,
        "strike": run.font.strike,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "font_color_rgb": run.font.color.rgb,
        "font_color_theme": run.font.color.theme_color,
    }


def apply_run_snapshot(run: Any, snapshot: Dict[str, Any], only_if_missing: bool = True) -> None:
    """Apply captured formatting to a run."""
    def assign(attr: str, value: Any) -> None:
        current = getattr(run, attr)
        if value is None:
            return
        if only_if_missing and current is not None:
            return
        setattr(run, attr, value)

    assign("bold", snapshot.get("bold"))
    assign("italic", snapshot.get("italic"))

    if snapshot.get("underline") is not None and (not only_if_missing or run.font.underline is None):
        run.font.underline = snapshot["underline"]
    if snapshot.get("strike") is not None and (not only_if_missing or run.font.strike is None):
        run.font.strike = snapshot["strike"]
    if snapshot.get("font_name") and (not only_if_missing or not run.font.name):
        run.font.name = snapshot["font_name"]
    if snapshot.get("font_size") and (not only_if_missing or not run.font.size):
        run.font.size = snapshot["font_size"]
    if snapshot.get("font_color_rgb") and (not only_if_missing or not run.font.color.rgb):
        run.font.color.rgb = snapshot["font_color_rgb"]
    elif snapshot.get("font_color_theme") and (not only_if_missing or not run.font.color.theme_color):
        run.font.color.theme_color = snapshot["font_color_theme"]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    """Remove run-like children from a paragraph but keep paragraph formatting."""
    for child in list(paragraph._p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in {"r", "hyperlink", "fldSimple"}:
            paragraph._p.remove(child)


def resolve_style(doc: DocxDocument, style_ids: List[str], style_names: List[str]) -> Optional[Any]:
    """Find a style by Word style_id first, then by display name."""
    for style in doc.styles:
        if getattr(style, "style_id", None) in style_ids:
            return style
    for style in doc.styles:
        if getattr(style, "name", None) in style_names:
            return style
    return None


def add_heading_paragraph(doc: DocxDocument, level: int) -> Paragraph:
    paragraph = doc.add_paragraph()
    heading_level = max(1, min(level, 6))
    style = resolve_style(
        doc,
        style_ids=[f"Heading{heading_level}"],
        style_names=[f"Heading {heading_level}", f"heading {heading_level}"],
    )
    if style is not None:
        paragraph.style = style
    return paragraph


def add_list_paragraph(doc: DocxDocument, is_ordered: bool, level: int) -> Paragraph:
    paragraph = doc.add_paragraph()
    normalized_level = min(level, 2)
    prefix = "ListNumber" if is_ordered else "ListBullet"
    style = resolve_style(
        doc,
        style_ids=[prefix if normalized_level == 0 else f"{prefix}{normalized_level + 1}"],
        style_names=[
            "List Number" if is_ordered else "List Bullet",
            f"List Number {normalized_level + 1}" if is_ordered else f"List Bullet {normalized_level + 1}",
        ],
    )
    if style is not None:
        paragraph.style = style
    return paragraph


def apply_quote_style(doc: DocxDocument, paragraph: Paragraph) -> None:
    style = resolve_style(doc, style_ids=["Quote"], style_names=["Quote"])
    if style is not None:
        paragraph.style = style


def parse_table(lines: List[str], start_index: int) -> Tuple[Optional[List[List[str]]], int]:
    table_lines: List[str] = []
    index = start_index

    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            index += 1
        else:
            break

    if len(table_lines) < 2:
        return None, start_index + 1

    table_data: List[List[str]] = []
    for line in table_lines:
        if "---" in line or ":-:" in line or ":--" in line or "--:" in line:
            continue
        table_data.append([cell.strip() for cell in line.split("|")[1:-1]])
    return table_data, index


def add_table_to_doc(table_data: List[List[str]], doc: DocxDocument) -> Optional[Table]:
    if not table_data:
        return None

    rows = len(table_data)
    columns = max(len(row) for row in table_data)
    word_table = doc.add_table(rows=rows, cols=columns)

    table_style = resolve_style(doc, style_ids=["TableGrid"], style_names=["Table Grid"])
    if table_style is not None:
        word_table.style = table_style

    for row_index, row_data in enumerate(table_data):
        for column_index, cell_text in enumerate(row_data):
            cell = word_table.cell(row_index, column_index)
            if cell.paragraphs:
                clear_paragraph_content(cell.paragraphs[0])
                parse_inline_formatting(cell_text, cell.paragraphs[0])
    return word_table


def process_list_items(
    lines: List[str],
    start_index: int,
    doc: DocxDocument,
    is_ordered: bool = False,
    level: int = 0,
    return_elements: bool = False,
) -> Tuple[int, Optional[List[Any]]]:
    elements: Optional[List[Any]] = [] if return_elements else None
    index = start_index

    while index < len(lines):
        stripped = lines[index].strip()
        original = lines[index]
        indent = len(original) - len(original.lstrip())
        current_level = indent // 3
        if current_level != level:
            break

        pattern = r"^\d+\.\s+(.+)" if is_ordered else r"^[-*+]\s+(.+)"
        match = re.match(pattern, stripped)
        if not match:
            break

        paragraph = add_list_paragraph(doc, is_ordered=is_ordered, level=level)
        parse_inline_formatting(match.group(1), paragraph)
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)

        index += 1

        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                index += 1
                continue
            next_original = lines[index]
            next_indent = len(next_original) - len(next_original.lstrip())
            next_level = next_indent // 3

            if next_level > level:
                nested_ordered = bool(re.match(r"^\d+\.\s+", next_line))
                nested_unordered = bool(re.match(r"^[-*+]\s+", next_line))
                if nested_ordered or nested_unordered:
                    index, nested = process_list_items(
                        lines,
                        index,
                        doc,
                        is_ordered=nested_ordered,
                        level=next_level,
                        return_elements=return_elements,
                    )
                    if return_elements and nested:
                        elements.extend(nested)
                else:
                    break
            else:
                break

    return index, elements


ORDERED_LIST_PATTERN = re.compile(r"^\d+\.\s+")
UNORDERED_LIST_PATTERN = re.compile(r"^[-*+]\s+")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
PAGE_BREAK_PATTERN = re.compile(r"^-{3,}\s*$")
HORIZONTAL_LINE_PATTERN = re.compile(r"^\*{3,}\s*$")
IMAGE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
TABLE_LINE_PATTERN = re.compile(r"^\|.+\|$")
QUOTE_PATTERN = re.compile(r"^>\s*(.+)$")

ALIGNMENT_MAP = {
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
}
ALIGN_INLINE_RE = re.compile(
    r'^(?:<center>(.*)</center>|<div\s+align="(right|center|justify|left)">(.*)</div>)$',
    re.IGNORECASE,
)
ALIGN_OPEN_RE = re.compile(
    r'^(?:<center>|<div\s+align="(right|center|justify|left)">)\s*$',
    re.IGNORECASE,
)
ALIGN_CLOSE_RE = re.compile(r"^</(?:center|div)>\s*$", re.IGNORECASE)

BLOCK_PATTERNS = [
    ORDERED_LIST_PATTERN,
    UNORDERED_LIST_PATTERN,
    HEADING_PATTERN,
    PAGE_BREAK_PATTERN,
    HORIZONTAL_LINE_PATTERN,
    IMAGE_PATTERN,
    TABLE_LINE_PATTERN,
    QUOTE_PATTERN,
]


def detect_alignment(line: str) -> Optional[Tuple[Optional[str], int]]:
    match = ALIGN_INLINE_RE.match(line)
    if match:
        if match.group(1) is not None:
            return match.group(1).strip(), WD_ALIGN_PARAGRAPH.CENTER
        return match.group(3).strip(), ALIGNMENT_MAP.get(match.group(2).lower(), WD_ALIGN_PARAGRAPH.LEFT)

    match = ALIGN_OPEN_RE.match(line)
    if match:
        return None, ALIGNMENT_MAP.get((match.group(1) or "center").lower(), WD_ALIGN_PARAGRAPH.CENTER)
    return None


def contains_block_markdown(value: str) -> bool:
    for line in value.split("\n"):
        stripped = line.strip()
        if any(pattern.match(stripped) for pattern in BLOCK_PATTERNS):
            return True
        if detect_alignment(stripped) is not None:
            return True
    return False


def add_horizontal_line(doc: DocxDocument) -> Paragraph:
    paragraph = doc.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)
    properties.append(border)
    return paragraph


def add_image_to_doc(doc: DocxDocument, url: str, alt_text: str, max_width_inches: Optional[float] = None) -> None:
    try:
        response = requests.get(url, timeout=20)
        response.raise_for_status()

        if max_width_inches is None:
            section = doc.sections[-1]
            max_width_inches = (section.page_width - section.left_margin - section.right_margin) / 914400

        image_stream = io.BytesIO(response.content)
        doc.add_picture(image_stream, width=Inches(max_width_inches))

        if alt_text:
            caption = doc.add_paragraph()
            caption.add_run(alt_text).italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as exc:
        logger.warning("Failed to add image from %s: %s", url, exc)
        doc.add_paragraph().add_run(f"[Image could not be loaded: {url}]")


def process_alignment_block(
    lines: List[str],
    start_index: int,
    doc: DocxDocument,
    alignment: int,
    return_elements: bool = False,
) -> Tuple[int, Optional[List[Any]]]:
    elements: Optional[List[Any]] = [] if return_elements else None
    index = start_index

    while index < len(lines):
        stripped = lines[index].strip()
        if ALIGN_CLOSE_RE.match(stripped):
            index += 1
            break
        if not stripped:
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = alignment
        parse_inline_formatting(stripped, paragraph)
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)
        index += 1

    return index, elements


def _add_field(paragraph: Paragraph, field_code: str) -> None:
    for field_type, text in [("begin", None), (None, field_code), ("end", None)]:
        run = paragraph.add_run()
        if field_type:
            field = OxmlElement("w:fldChar")
            field.set(qn("w:fldCharType"), field_type)
            run._r.append(field)
        else:
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = f" {text} "
            run._r.append(instruction)


PAGE_TOKEN_RE = re.compile(r"(\{page}|\{pages})")


def set_header_footer(doc: DocxDocument, text: str, kind: str = "header") -> None:
    token_map = {"{page}": "PAGE", "{pages}": "NUMPAGES"}

    def fill_paragraph(paragraph: Paragraph, content: str) -> None:
        existing_alignment = paragraph.alignment
        clear_paragraph_content(paragraph)
        for part in PAGE_TOKEN_RE.split(content):
            if part in token_map:
                _add_field(paragraph, token_map[part])
            elif part:
                paragraph.add_run(part)
        paragraph.alignment = existing_alignment if existing_alignment is not None else WD_ALIGN_PARAGRAPH.CENTER

    def update_part(section_part: Any) -> None:
        section_part.is_linked_to_previous = False
        if section_part.paragraphs:
            fill_paragraph(section_part.paragraphs[0], text)
        else:
            fill_paragraph(section_part.add_paragraph(), text)

    for section in doc.sections:
        update_part(getattr(section, kind))
        if section.different_first_page_header_footer:
            first_part = getattr(section, f"first_page_{kind}", None)
            if first_part is not None:
                update_part(first_part)
        even_part = getattr(section, f"even_page_{kind}", None)
        if even_part is not None and doc.settings.element.find(qn("w:evenAndOddHeaders")) is not None:
            update_part(even_part)


def add_toc(doc: DocxDocument) -> None:
    heading = add_heading_paragraph(doc, level=1)
    heading.add_run("Table of Contents")
    paragraph = doc.add_paragraph()

    begin = paragraph.add_run()
    begin_field = OxmlElement("w:fldChar")
    begin_field.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_field)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    instruction_run._r.append(instruction)

    separate = paragraph.add_run()
    separate_field = OxmlElement("w:fldChar")
    separate_field.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_field)

    paragraph.add_run("[Table of Contents - open in Word and press F9 to update]")

    end = paragraph.add_run()
    end_field = OxmlElement("w:fldChar")
    end_field.set(qn("w:fldCharType"), "end")
    end._r.append(end_field)

    doc.add_page_break()

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)


def process_markdown_block(doc: DocxDocument, lines: List[str], start_index: int, return_element: bool = True) -> Tuple[int, List[Any]]:
    stripped = lines[start_index].strip()
    elements: List[Any] = []

    def collect(element: Any) -> None:
        if return_element:
            elements.append(element)
            doc._body._body.remove(element)

    try:
        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            heading = add_heading_paragraph(doc, level=len(heading_match.group(1)))
            parse_inline_formatting(heading_match.group(2), heading)
            collect(heading._p)
            return start_index + 1, elements

        if TABLE_LINE_PATTERN.match(stripped):
            table_data, next_index = parse_table(lines, start_index)
            if table_data:
                word_table = add_table_to_doc(table_data, doc)
                if word_table is not None:
                    collect(word_table._tbl)
                return next_index, elements

        quote_match = QUOTE_PATTERN.match(stripped)
        if quote_match:
            paragraph = doc.add_paragraph()
            apply_quote_style(doc, paragraph)
            parse_inline_formatting(quote_match.group(1), paragraph)
            collect(paragraph._p)
            return start_index + 1, elements

        if PAGE_BREAK_PATTERN.match(stripped):
            doc.add_page_break()
            collect(doc.paragraphs[-1]._p)
            return start_index + 1, elements

        if HORIZONTAL_LINE_PATTERN.match(stripped):
            collect(add_horizontal_line(doc)._p)
            return start_index + 1, elements

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            add_image_to_doc(doc, image_match.group(2), image_match.group(1))
            return start_index + 1, elements

        align_result = detect_alignment(stripped)
        if align_result is not None:
            inner, alignment = align_result
            if inner is not None:
                paragraph = doc.add_paragraph()
                paragraph.alignment = alignment
                parse_inline_formatting(inner, paragraph)
                collect(paragraph._p)
                return start_index + 1, elements
            next_index, block_elements = process_alignment_block(lines, start_index + 1, doc, alignment, return_elements=return_element)
            if return_element and block_elements:
                elements.extend(block_elements)
            return next_index, elements

        if ORDERED_LIST_PATTERN.match(stripped):
            next_index, list_elements = process_list_items(lines, start_index, doc, is_ordered=True, level=0, return_elements=return_element)
            return next_index, list_elements or []

        if UNORDERED_LIST_PATTERN.match(stripped):
            next_index, list_elements = process_list_items(lines, start_index, doc, is_ordered=False, level=0, return_elements=return_element)
            return next_index, list_elements or []

        paragraph = doc.add_paragraph()
        parse_inline_formatting(stripped, paragraph)
        collect(paragraph._p)
        return start_index + 1, elements
    except Exception as exc:
        logger.error("Failed to process markdown block at line %d: %s", start_index, exc, exc_info=True)
        return start_index + 1, elements


def render_markdown_into_document(doc: DocxDocument, markdown_content: str) -> Dict[str, int]:
    """Append markdown content to a document."""
    lines = markdown_content.split("\n")
    index = 0
    stats = {
        "headings": 0,
        "tables": 0,
        "ordered_lists": 0,
        "unordered_lists": 0,
        "paragraphs": 0,
    }

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()
        if not stripped:
            index += 1
            continue

        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            heading = add_heading_paragraph(doc, level=len(heading_match.group(1)))
            parse_inline_formatting(heading_match.group(2), heading)
            stats["headings"] += 1
            index += 1
            continue

        if TABLE_LINE_PATTERN.match(stripped):
            table_data, next_index = parse_table(lines, index)
            if table_data:
                add_table_to_doc(table_data, doc)
                stats["tables"] += 1
                index = next_index
                continue

        quote_match = QUOTE_PATTERN.match(stripped)
        if quote_match:
            paragraph = doc.add_paragraph()
            apply_quote_style(doc, paragraph)
            parse_inline_formatting(quote_match.group(1), paragraph)
            stats["paragraphs"] += 1
            index += 1
            continue

        if PAGE_BREAK_PATTERN.match(stripped):
            doc.add_page_break()
            index += 1
            continue

        if HORIZONTAL_LINE_PATTERN.match(stripped):
            add_horizontal_line(doc)
            index += 1
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            add_image_to_doc(doc, image_match.group(2), image_match.group(1))
            index += 1
            continue

        align_result = detect_alignment(stripped)
        if align_result is not None:
            inner, alignment = align_result
            if inner is not None:
                paragraph = doc.add_paragraph()
                paragraph.alignment = alignment
                parse_inline_formatting(inner, paragraph)
                stats["paragraphs"] += 1
                index += 1
                continue
            index, _ = process_alignment_block(lines, index + 1, doc, alignment, return_elements=False)
            stats["paragraphs"] += 1
            continue

        if ORDERED_LIST_PATTERN.match(stripped):
            index, _ = process_list_items(lines, index, doc, is_ordered=True, level=0, return_elements=False)
            stats["ordered_lists"] += 1
            continue

        if UNORDERED_LIST_PATTERN.match(stripped):
            index, _ = process_list_items(lines, index, doc, is_ordered=False, level=0, return_elements=False)
            stats["unordered_lists"] += 1
            continue

        paragraph = doc.add_paragraph()
        parse_inline_formatting(raw_line.rstrip(), paragraph)
        stats["paragraphs"] += 1
        index += 1

    return stats


def insert_markdown_content_after_paragraph(doc: DocxDocument, paragraph: Paragraph, content: str) -> None:
    body = doc._body._body
    paragraph_index = list(body).index(paragraph._p)
    inserted = 0
    lines = content.split("\n")
    index = 0

    while index < len(lines):
        if not lines[index].strip():
            index += 1
            continue
        index, elements = process_markdown_block(doc, lines, index, return_element=True)
        for element in elements:
            body.insert(paragraph_index + 1 + inserted, element)
            inserted += 1


def _replace_single_placeholder_in_paragraph(
    paragraph: Paragraph,
    placeholder: str,
    value: str,
    doc: Optional[DocxDocument] = None,
) -> bool:
    full_text = paragraph.text
    if placeholder not in full_text:
        return False

    runs = list(paragraph.runs)
    if not runs:
        return False

    combined = ""
    run_map: List[Tuple[int, int, Any]] = []
    for run in runs:
        start = len(combined)
        combined += run.text
        run_map.append((start, len(combined), run))

    placeholder_start = combined.find(placeholder)
    if placeholder_start == -1:
        return False
    placeholder_end = placeholder_start + len(placeholder)

    formatting_run = None
    for start, end, run in run_map:
        if start <= placeholder_start < end:
            formatting_run = run
            break

    snapshot = snapshot_run_format(formatting_run) if formatting_run is not None else {}
    text_before = combined[:placeholder_start]
    text_after = combined[placeholder_end:]
    has_block_content = contains_block_markdown(value)

    clear_paragraph_content(paragraph)

    if text_before:
        before_run = paragraph.add_run(text_before)
        apply_run_snapshot(before_run, snapshot)

    if has_block_content and doc is not None:
        if text_after:
            after_run = paragraph.add_run(text_after)
            apply_run_snapshot(after_run, snapshot)
        insert_markdown_content_after_paragraph(doc, paragraph, value)
        return True

    existing_runs = len(paragraph.runs)
    parse_inline_formatting(value, paragraph)
    for run in list(paragraph.runs)[existing_runs:]:
        apply_run_snapshot(run, snapshot, only_if_missing=True)

    if text_after:
        after_run = paragraph.add_run(text_after)
        apply_run_snapshot(after_run, snapshot)

    return True


def replace_placeholders_in_paragraph(paragraph: Paragraph, context: Dict[str, str], doc: Optional[DocxDocument] = None) -> None:
    iterations = 0
    while iterations < 100:
        iterations += 1
        matches = PLACEHOLDER_PATTERN.findall(paragraph.text)
        if not matches:
            return

        replaced = False
        for placeholder_name in matches:
            if placeholder_name not in context:
                continue
            replacement = "" if context[placeholder_name] is None else str(context[placeholder_name])
            for placeholder in [f"{{{{{{{placeholder_name}}}}}}}", f"{{{{{placeholder_name}}}}}"]:
                if placeholder in paragraph.text and _replace_single_placeholder_in_paragraph(paragraph, placeholder, replacement, doc=doc):
                    replaced = True
                    break
            if replaced:
                break
        if not replaced:
            return


def replace_placeholders_in_table(table: Table, context: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in list(cell.paragraphs):
                replace_placeholders_in_paragraph(paragraph, context, doc=None)


def iter_section_parts(section: Any) -> Iterable[Any]:
    parts = []
    seen = set()
    for attribute in [
        "header",
        "footer",
        "first_page_header",
        "first_page_footer",
        "even_page_header",
        "even_page_footer",
    ]:
        part = getattr(section, attribute, None)
        if part is None:
            continue
        key = getattr(getattr(part, "_element", None), "part", None) or id(part)
        if key in seen:
            continue
        seen.add(key)
        parts.append(part)
    return parts


def replace_placeholders_in_document(doc: DocxDocument, context: Dict[str, str]) -> None:
    for paragraph in list(doc.paragraphs):
        replace_placeholders_in_paragraph(paragraph, context, doc=doc)

    for table in doc.tables:
        replace_placeholders_in_table(table, context)

    for section in doc.sections:
        for part in iter_section_parts(section):
            for paragraph in list(part.paragraphs):
                replace_placeholders_in_paragraph(paragraph, context, doc=None)
            for table in part.tables:
                replace_placeholders_in_table(table, context)


def set_document_metadata(doc: DocxDocument, title: Optional[str] = None, author: Optional[str] = None, subject: Optional[str] = None) -> None:
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author
    if subject:
        doc.core_properties.subject = subject


def create_document(template_path: Optional[str] = None) -> DocxDocument:
    return Document(template_path) if template_path else Document()
